from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import io
import os
import tempfile
import zipfile
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
import base64

app = FastAPI(title="PDF Toolkit API")

# Add CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"]
)

# --- Helper Functions (Using ONLY fitz/PyMuPDF) ---

def merge_pdfs_logic(pdf_bytes_list):
    try:
        merged_doc = fitz.open()
        for pdf_bytes in pdf_bytes_list:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            merged_doc.insert_pdf(doc)
            doc.close()
        
        output = io.BytesIO(merged_doc.tobytes())
        merged_doc.close()
        return output
    except Exception as e:
        raise Exception(f"Merge error: {str(e)}")

def split_pdf_logic(pdf_bytes, range_str):
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        outputs = []
        
        # Parse range string "1-3, 5" -> [(1,3), (5,5)]
        ranges = []
        for r in range_str.split(','):
            r = r.strip()
            if '-' in r:
                start, end = map(int, r.split('-'))
                ranges.append((start, end))
            else:
                page = int(r)
                ranges.append((page, page))
                
        for page_range in ranges:
            start, end = page_range
            new_doc = fitz.open()
            # Adjust for 0-based index
            # fitz uses 0-based indexing
            # User input is 1-based
            
            # Create list of page numbers to include
            # range(start-1, end) because python range is exclusive at end
            # but user says "1-3" meaning 1,2,3. So index 0,1,2.
            # Python range(0, 3) gives 0,1,2.
            # So range(start-1, end) is correct.
            
            final_end = min(end, len(doc))
            if start - 1 < len(doc):  # Check valid start
                new_doc.insert_pdf(doc, from_page=start-1, to_page=final_end-1)
            
            output = io.BytesIO(new_doc.tobytes())
            new_doc.close()
            outputs.append(output)
            
        return outputs
    except Exception as e:
        raise Exception(f"Split error: {str(e)}")

def validate_pdf(pdf_bytes):
    """Validate if PDF is valid and not empty"""
    try:
        if len(pdf_bytes) < 100:  # Too small to be valid
            return False
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if len(doc) == 0:
            return False
        doc.close()
        return True
    except:
        return False

def compress_pdf_logic(pdf_bytes, level="medium"):
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        if level == "low":
            image_quality = 85; dpi = 150
        elif level == "medium":
            image_quality = 60; dpi = 120
        else: # high
            image_quality = 40; dpi = 90
            
        for page in doc:
            for img in page.get_images(full=True):
                xref = img[0]
                try:
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image = Image.open(io.BytesIO(image_bytes))
                    
                    width, height = image.size
                    new_width = int(width * (dpi / 150))
                    new_height = int(height * (dpi / 150))
                    
                    if new_width > 0 and new_height > 0:
                        image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
                        
                    img_output = io.BytesIO()
                    if image.mode not in ('RGB', 'L'):
                        image = image.convert('RGB')
                        
                    image.save(img_output, format='JPEG', quality=image_quality, optimize=True)
                    page.replace_image(xref, stream=img_output.getvalue())
                except:
                    continue
                    
        output = io.BytesIO()
        # Use less aggressive garbage collection to prevent corruption
        doc.save(
            output,
            garbage=3,
            deflate=True,
            clean=True,
            deflate_images=True,
            deflate_fonts=True
        )
        doc.close()
        
        # Validate output
        if not validate_pdf(output.getvalue()):
            # If compression failed, return original
            return io.BytesIO(pdf_bytes)
            
        output.seek(0)
        return output
    except Exception as e:
        raise Exception(f"Compress error: {str(e)}")

def pdf_to_word_logic(pdf_bytes):
    try:
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        doc = Document()
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            text = page.get_text()
            if page_num > 0: doc.add_page_break()
            
            heading = doc.add_paragraph(f"Page {page_num + 1}")
            heading.style = 'Heading 2'
            
            if text.strip():
                for para_text in text.split('\n\n'):
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    except Exception as e:
        raise Exception(f"Conversion error: {str(e)}")

def word_to_pdf_logic(docx_bytes):
    try:
        doc = Document(io.BytesIO(docx_bytes))
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
            tmp_path = tmp_pdf.name
            
        pdf_doc = SimpleDocTemplate(tmp_path, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        for para in doc.paragraphs:
            if para.text.strip():
                style = styles['Heading1'] if para.style.name.startswith('Heading') else styles['Normal']
                story.append(Paragraph(para.text, style))
                story.append(Spacer(1, 12))
                
        pdf_doc.build(story)
        
        with open(tmp_path, 'rb') as f:
            output = io.BytesIO(f.read())
        os.unlink(tmp_path)
        return output
    except Exception as e:
        raise Exception(f"Word to PDF error: {str(e)}")

def rotate_pdf_logic(pdf_bytes, angle, pages_to_rotate):
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        angle = int(angle)
        target_pages = []
        if pages_to_rotate != "all":
            # "1, 3, 5" -> [1, 3, 5]
            target_pages = [int(p.strip()) for p in pages_to_rotate.split(',')]
            
        for i, page in enumerate(doc):
            if pages_to_rotate == "all" or (i + 1) in target_pages:
                page.set_rotation(angle)
                
        output = io.BytesIO(doc.tobytes())
        return output
    except Exception as e:
        raise Exception(f"Rotation error: {str(e)}")

# --- Endpoints ---

@app.post("/merge")
async def merge_pdfs_endpoint(files: list[UploadFile]):
    try:
        file_bytes_list = [await f.read() for f in files]
        result = merge_pdfs_logic(file_bytes_list)
        return StreamingResponse(
            result, 
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=merged.pdf"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/split")
async def split_pdf_endpoint(file: UploadFile, ranges: str = Form(...)):
    try:
        content = await file.read()
        split_files = split_pdf_logic(content, ranges)
        
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
            for i, pdf_io in enumerate(split_files):
                zip_file.writestr(f"split_part_{i+1}.pdf", pdf_io.getvalue())
                
        zip_buffer.seek(0)
        return StreamingResponse(
            zip_buffer,
            media_type="application/zip",
            headers={"Content-Disposition": "attachment; filename=split_files.zip"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/compress")
async def compress_pdf_endpoint(file: UploadFile, level: str = Form("medium")):
    try:
        content = await file.read()
        result = compress_pdf_logic(content, level)
        return StreamingResponse(
            result, 
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=compressed.pdf"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/pdf-to-word")
async def pdf_to_word_endpoint(file: UploadFile):
    try:
        content = await file.read()
        result = pdf_to_word_logic(content)
        return StreamingResponse(
            result, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=converted.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/word-to-pdf")
async def word_to_pdf_endpoint(file: UploadFile):
    try:
        content = await file.read()
        result = word_to_pdf_logic(content)
        return StreamingResponse(
            result, 
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=converted.pdf"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/rotate")
async def rotate_pdf_endpoint(file: UploadFile, angle: int = Form(...), pages: str = Form("all")):
    try:
        content = await file.read()
        result = rotate_pdf_logic(content, angle, pages)
        return StreamingResponse(
            result, 
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=rotated.pdf"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/preview")
async def preview_pdf_endpoint(file: UploadFile, limit: int = Form(default=5)):
    """Returns list of base64 images. Use limit=0 for all pages."""
    try:
        content = await file.read()
        doc = fitz.open(stream=content, filetype="pdf")
        
        images = []
        max_pages = len(doc)
        if limit > 0:
            max_pages = min(limit, max_pages)

        for i in range(max_pages):
            page = doc[i]
            # Matrix(2.0, 2.0) = Roughly 144 DPI (Double resolution)
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
            img_data = pix.tobytes("png")
            b64_img = base64.b64encode(img_data).decode('utf-8')
            images.append(f"data:image/png;base64,{b64_img}")
            
        return {"pages": images, "total_pages": len(doc)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

def reorder_pdf_logic(pdf_bytes, page_order_str):
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        new_doc = fitz.open()
        
        # Parse "1, 3, 2" -> [0, 2, 1] (0-indexed)
        try:
            page_order = [int(x.strip()) - 1 for x in page_order_str.split(',')]
        except ValueError:
            raise Exception("Invalid page order format. Use comma-separated numbers.")
            
        for page_idx in page_order:
            if 0 <= page_idx < len(doc):
                new_doc.insert_pdf(doc, from_page=page_idx, to_page=page_idx)
            
        output = io.BytesIO(new_doc.tobytes())
        return output
    except Exception as e:
        raise Exception(f"Reorder error: {str(e)}")

def assemble_pdf_logic(pdf_bytes_list, manifest):
    # manifest = [{"file_index": 0, "page_index": 0, "rotation": 0}, ...]
    try:
        source_docs = []
        for pdf_bytes in pdf_bytes_list:
            source_docs.append(fitz.open(stream=pdf_bytes, filetype="pdf"))
            
        new_doc = fitz.open()
        
        for item in manifest:
            f_idx = item.get("file_index")
            p_idx = item.get("page_index")
            rot = item.get("rotation", 0)
            
            if 0 <= f_idx < len(source_docs):
                src_doc = source_docs[f_idx]
                if 0 <= p_idx < len(src_doc):
                    # Insert page
                    new_doc.insert_pdf(src_doc, from_page=p_idx, to_page=p_idx)
                    # Rotate if needed (relative to current rotation)
                    if rot != 0:
                        # fitz set_rotation is absolute. 
                        # But user usually expects relative add. 
                        # Let's assume input 'rotation' is the DESIRED final rotation 
                        # or relative? Adobe usually rotates relative to view.
                        # Let's assume the manifest sends the Absolute rotation desired.
                        # Actually simple: just rotate the last added page.
                        new_doc[-1].set_rotation(rot)

        output = io.BytesIO(new_doc.tobytes())
        return output
    except Exception as e:
        raise Exception(f"Assembly error: {str(e)}")

def watermark_pdf_logic(pdf_bytes, text, opacity=0.3, rotation=45, color=(0, 0, 0), scale=1.0):
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        for page in doc:
            page_width = page.rect.width
            page_height = page.rect.height
            
            # Dynamic Font Size: ~15% of page width
            # This ensures consistent look regardless of PDF resolution
            base_font_size = page_width * 0.15
            font_size = int(base_font_size)
            
            # Create a PIL image for the watermark
            text_length = len(text)
            # Base dimensions
            # We estimate text width roughly. 
            # Better to measure if we had the font loaded, but estimation is okay for container size.
            img_width = int(text_length * font_size * 0.6 * scale)
            img_height = int(font_size * 2 * scale)
            
            # Create transparent image
            watermark_img = Image.new('RGBA', (img_width, img_height), (255, 255, 255, 0))
            draw = ImageDraw.Draw(watermark_img)
            
            # Try to use a font, fallback to reasonable system fonts
            font_paths = [
                "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", # Docker/Linux
                "C:/Windows/Fonts/arial.ttf", # Windows
                "C:/Windows/Fonts/Arial.ttf",
                "Arial.ttf",
                "/System/Library/Fonts/Helvetica.ttc" # Mac
            ]
            
            font = None
            scaled_size = int(font_size * scale)
            
            for path in font_paths:
                try:
                    font = ImageFont.truetype(path, scaled_size)
                    break 
                except:
                    continue
            
            if font is None:
                try:
                    # Last resort: load_default (might be tiny on some systems)
                    font = ImageFont.load_default()
                    # Warn or log?
                    print("Warning: Using default PIL font (not scalable)")
                except:
                    pass
            
            # Calculate text color with opacity
            alpha = int(opacity * 255)
            text_color = (128, 128, 128, alpha)  # Gray with alpha
            
            # Draw text
            # If font is loaded, use it. If not, default might not support size.
            # For simplicity in container, we try to draw.
            draw.text((10, 10), text, fill=text_color, font=font)
            
            # Rotate the image
            watermark_img = watermark_img.rotate(rotation, expand=True)
            
            # Save to bytes
            img_byte_arr = io.BytesIO()
            watermark_img.save(img_byte_arr, format='PNG')
            img_bytes = img_byte_arr.getvalue()
            
            # Center position
            x = (page_width - watermark_img.width) / 2
            y = (page_height - watermark_img.height) / 2
            
            # Insert image
            img_rect = fitz.Rect(x, y, x + watermark_img.width, y + watermark_img.height)
            page.insert_image(img_rect, stream=img_bytes, overlay=True)
            
        output = io.BytesIO(doc.tobytes())
        return output
    except Exception as e:
        raise Exception(f"Watermark error: {str(e)}")

def watermark_pdf_logic_enhanced(pdf_bytes, text=None, image_bytes=None, opacity=0.3, rotation=45, scale=1.0, color_mode="original"):
    """Enhanced watermark function - designed to EXACTLY match the frontend preview in PDFEditor.jsx.
    
    Frontend preview formula (at zoomLevel=1):
    - Text fontSize = 800 * watermarkScale * 0.15 = 120px at scale 1.0 (15% of 800px width)
    - Image maxWidth = 300 * watermarkScale, maxHeight = 200 * watermarkScale
    
    For PDF, we maintain the SAME RATIO relative to page width:
    - Text should be 15% of page width at scale 1.0
    - Image should be proportionally scaled
    """
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        for page in doc:
            page_rect = page.rect
            page_width = page_rect.width
            page_height = page_rect.height
            
            # Center of the page
            center_x = page_rect.x0 + page_width / 2
            center_y = page_rect.y0 + page_height / 2
            
            # Process image watermark if provided
            if image_bytes:
                try:
                    img = Image.open(io.BytesIO(image_bytes))
                    
                    if img.mode != 'RGBA':
                        img = img.convert('RGBA')
                    
                    # Apply color mode transformation
                    if color_mode == "grayscale":
                        r, g, b, a = img.split()
                        gray = img.convert('L')
                        img = Image.merge('RGBA', (gray, gray, gray, a))
                    elif color_mode == "bw":
                        r, g, b, a = img.split()
                        gray = img.convert('L')
                        bw = gray.point(lambda x: 0 if x < 128 else 255, 'L')
                        img = Image.merge('RGBA', (bw, bw, bw, a))
                    
                    # Apply opacity
                    r, g, b, a = img.split()
                    a = a.point(lambda x: int(x * opacity))
                    img = Image.merge('RGBA', (r, g, b, a))
                    
                    # Frontend uses 300x200 on 800px preview = 37.5% x 25% of preview
                    # Apply same ratio to PDF page
                    max_width = page_width * 0.375 * scale
                    max_height = page_height * 0.25 * scale
                    
                    img_ratio = img.width / img.height if img.height > 0 else 1
                    if img.width > max_width or img.height > max_height:
                        if img_ratio > max_width / max_height:
                            new_width = int(max_width)
                            new_height = int(max_width / img_ratio)
                        else:
                            new_height = int(max_height)
                            new_width = int(max_height * img_ratio)
                    else:
                        # Scale smaller images up to fit
                        scale_up = min(max_width / img.width, max_height / img.height)
                        new_width = int(img.width * scale_up)
                        new_height = int(img.height * scale_up)
                    
                    img = img.resize((max(1, new_width), max(1, new_height)), Image.Resampling.LANCZOS)
                    
                    # Rotate - CSS uses clockwise for positive angles, PIL uses counter-clockwise
                    # So we negate the angle to match frontend preview
                    if rotation != 0:
                        img = img.rotate(-rotation, expand=True, resample=Image.Resampling.BICUBIC)
                    
                    img_byte_arr = io.BytesIO()
                    img.save(img_byte_arr, format='PNG')
                    img_data = img_byte_arr.getvalue()
                    
                    img_x = center_x - img.width / 2
                    img_y = center_y - img.height / 2
                    if text:
                        img_y -= page_height * 0.03  # Offset up slightly
                    
                    img_rect = fitz.Rect(img_x, img_y, img_x + img.width, img_y + img.height)
                    page.insert_image(img_rect, stream=img_data, overlay=True)
                    
                except Exception as img_error:
                    print(f"Image watermark error: {str(img_error)}")
            
            # Process text watermark
            if text:
                # Frontend: fontSize = 800 * scale * 0.15 = 15% of 800px = 120px at scale 1.0
                # For PDF: fontSize = 15% of page_width at scale 1.0
                # This ensures the text appears the SAME relative size
                font_size = int(page_width * 0.15 * scale)
                font_size = max(24, font_size)  # Minimum readable size
                
                # Create a larger canvas to ensure text fits after rotation
                # Width estimate: each character ~0.6 of font size, plus padding
                text_length = len(text)
                canvas_width = int(text_length * font_size * 0.8) + 100
                canvas_height = int(font_size * 2) + 50
                
                watermark_img = Image.new('RGBA', (canvas_width, canvas_height), (255, 255, 255, 0))
                draw = ImageDraw.Draw(watermark_img)
                
                # Load font
                font_paths = [
                    "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
                    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
                    "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
                    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                    "C:/Windows/Fonts/arialbd.ttf",
                    "C:/Windows/Fonts/arial.ttf",
                    "/System/Library/Fonts/Helvetica.ttc"
                ]
                
                font = None
                for path in font_paths:
                    try:
                        font = ImageFont.truetype(path, font_size)
                        break 
                    except:
                        continue
                
                if font is None:
                    try:
                        font = ImageFont.load_default()
                    except:
                        pass
                
                # Frontend uses rgba(60, 60, 60, opacity) 
                alpha = int(opacity * 255)
                text_color = (60, 60, 60, alpha)
                
                # Draw text centered in canvas
                text_x_pos = 50
                text_y_pos = 25
                draw.text((text_x_pos, text_y_pos), text, fill=text_color, font=font)
                
                # Crop to actual content to avoid huge empty space
                bbox = watermark_img.getbbox()
                if bbox:
                    watermark_img = watermark_img.crop(bbox)
                
                # Rotate - CSS uses clockwise for positive angles, PIL uses counter-clockwise
                # So we negate the angle to match frontend preview
                if rotation != 0:
                    watermark_img = watermark_img.rotate(-rotation, expand=True, resample=Image.Resampling.BICUBIC)
                
                text_byte_arr = io.BytesIO()
                watermark_img.save(text_byte_arr, format='PNG')
                text_data = text_byte_arr.getvalue()
                
                # Center on page
                text_x = center_x - watermark_img.width / 2
                text_y = center_y - watermark_img.height / 2
                if image_bytes:
                    text_y += page_height * 0.05  # Offset down if image above
                
                text_rect = fitz.Rect(text_x, text_y, text_x + watermark_img.width, text_y + watermark_img.height)
                page.insert_image(text_rect, stream=text_data, overlay=True)
        
        output = io.BytesIO(doc.tobytes())
        return output
    except Exception as e:
        raise Exception(f"Enhanced watermark error: {str(e)}")

@app.post("/reorder")
async def reorder_pdf_endpoint(file: UploadFile, order: str = Form(...)):
    try:
        content = await file.read()
        result = reorder_pdf_logic(content, order)
        return StreamingResponse(
            result, 
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=reordered.pdf"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/watermark")
async def watermark_pdf_endpoint(
    file: UploadFile = File(...), 
    text: str = Form(None),  # Now optional
    opacity: float = Form(0.3),
    rotation: int = Form(45),
    scale: float = Form(1.0),
    watermark_type: str = Form("text"),  # 'text', 'image', or 'both'
    color_mode: str = Form("original"),  # 'original', 'grayscale', or 'bw'
    image: UploadFile = File(None)  # Optional image file
):
    try:
        print(f"Watermark request: type={watermark_type}, text={text}, color_mode={color_mode}, scale={scale}")
        content = await file.read()
        
        # Read image if provided
        image_bytes = None
        if image and image.filename and (watermark_type == "image" or watermark_type == "both"):
            print(f"Reading image: {image.filename}")
            image_bytes = await image.read()
        
        # Determine what to apply based on watermark_type
        apply_text = watermark_type in ("text", "both") and text
        apply_image = watermark_type in ("image", "both") and image_bytes
        
        print(f"apply_text={apply_text}, apply_image={apply_image}")
        
        if not apply_text and not apply_image:
            raise HTTPException(status_code=400, detail="No watermark content provided. Please provide text or upload an image.")
        
        # Validation
        print(f"File size: {len(content)} bytes")
        print(f"File header: {content[:20]}")
        
        if not content.startswith(b'%PDF-'):
             print("Error: File does not start with %PDF-")
             # Try to be helpful: if it looks like an image, suggest converting
             from PIL import Image
             try:
                 img = Image.open(io.BytesIO(content))
                 raise HTTPException(status_code=400, detail=f"The uploaded file is an Image ({img.format}), not a PDF. Please convert it to PDF first.")
             except:
                 pass
             raise HTTPException(status_code=400, detail="The uploaded file is not a valid PDF.")

        if not validate_pdf(content):
             raise HTTPException(status_code=400, detail="PDF Validation failed (file may be empty or corrupted).")

        result = watermark_pdf_logic_enhanced(
            content, 
            text=text if apply_text else None,
            image_bytes=image_bytes if apply_image else None,
            opacity=opacity, 
            rotation=rotation, 
            scale=scale,
            color_mode=color_mode
        )
        return StreamingResponse(
            result, 
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=watermarked.pdf"}
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Watermark error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


def int_to_roman(num):
    """Convert integer to roman numeral"""
    val = [
        1000, 900, 500, 400,
        100, 90, 50, 40,
        10, 9, 5, 4,
        1
    ]
    syms = [
        'm', 'cm', 'd', 'cd',
        'c', 'xc', 'l', 'xl',
        'x', 'ix', 'v', 'iv',
        'i'
    ]
    roman_num = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syms[i]
            num -= val[i]
        i += 1
    return roman_num


def add_page_numbers_logic(pdf_bytes, format_type="number", position="bottom-center", 
                           start_from=1, font_size=12, color="#000000", margin=36, total_pages=None):
    """Add page numbers to PDF"""
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        total = len(doc) if total_pages is None else total_pages
        
        # Parse hex color to RGB tuple (0-1 range for fitz)
        hex_color = color.lstrip('#')
        r = int(hex_color[0:2], 16) / 255
        g = int(hex_color[2:4], 16) / 255
        b = int(hex_color[4:6], 16) / 255
        text_color = (r, g, b)
        
        for page_num, page in enumerate(doc):
            page_width = page.rect.width
            page_height = page.rect.height
            
            # Calculate actual page number based on start_from
            actual_num = page_num + start_from
            
            # Generate page number text based on format
            if format_type == "page-of":
                page_text = f"Page {actual_num} of {total + start_from - 1}"
            elif format_type == "roman":
                page_text = int_to_roman(actual_num)
            else:  # number
                page_text = str(actual_num)
            
            # Calculate accurate text width using fitz
            # Use get_text_length for proper measurement
            try:
                text_width = fitz.get_text_length(page_text, fontname="helv", fontsize=font_size)
            except:
                # Fallback: more accurate approximation (0.6 is average char width ratio for Helvetica)
                text_width = len(page_text) * font_size * 0.6
            
            # Y position
            if position.startswith("top"):
                y_top = margin
                y_bottom = margin + font_size + 10
            else:  # bottom
                y_top = page_height - margin - font_size - 10
                y_bottom = page_height - margin
            
            # X position and alignment
            if position.endswith("left"):
                x_left = margin
                x_right = margin + 150  # wide enough for any page number text
                align = fitz.TEXT_ALIGN_LEFT
            elif position.endswith("right"):
                x_left = page_width - 50 - 150  # 50pt from right edge, 150pt wide box
                x_right = page_width - 50
                align = fitz.TEXT_ALIGN_RIGHT
            else:  # center
                x_left = page_width / 4
                x_right = page_width * 3 / 4
                align = fitz.TEXT_ALIGN_CENTER
            
            # Create text rectangle
            text_rect = fitz.Rect(x_left, y_top, x_right, y_bottom)
            
            # Insert text in box with alignment
            page.insert_textbox(
                text_rect,
                page_text,
                fontsize=font_size,
                color=text_color,
                fontname="helv",
                align=align
            )
        
        output = io.BytesIO(doc.tobytes())
        doc.close()
        return output
    except Exception as e:
        raise Exception(f"Page numbering error: {str(e)}")


@app.post("/add_page_numbers")
async def add_page_numbers_endpoint(
    file: UploadFile,
    format: str = Form("number"),  # "number", "page-of", "roman"
    position: str = Form("bottom-center"),  # "top-left", "top-center", "top-right", "bottom-left", "bottom-center", "bottom-right"
    start_from: int = Form(1),
    font_size: int = Form(12),
    color: str = Form("#000000"),
    margin: int = Form(36),
    total_pages: int = Form(None)
):
    try:
        content = await file.read()
        result = add_page_numbers_logic(
            content, 
            format_type=format,
            position=position,
            start_from=start_from,
            font_size=font_size,
            color=color,
            margin=margin,
            total_pages=total_pages
        )
        return StreamingResponse(
            result, 
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=numbered.pdf"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/assemble")
async def assemble_pdf_endpoint(files: list[UploadFile], manifest: str = Form(...)):
    # manifest is a JSON string
    import json
    try:
        manifest_data = json.loads(manifest)
        file_bytes_list = [await f.read() for f in files]
        result = assemble_pdf_logic(file_bytes_list, manifest_data)
        return StreamingResponse(
            result, 
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=assembled.pdf"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/preview")
async def preview_pdf_endpoint(file: UploadFile, limit: int = Form(5)):
    """Returns list of base64 images. Set limit param to 0 or -1 for all pages."""
    try:
        content = await file.read()
        doc = fitz.open(stream=content, filetype="pdf")
        
        images = []
        
        page_range = range(len(doc))
        if limit > 0:
            page_range = range(min(limit, len(doc)))
            
        for i in page_range:
            page = doc[i]
            # Use ultra high res (was 3.0)
            # 5.0 is approx 360 dpi, print quality
            pix = page.get_pixmap(matrix=fitz.Matrix(5.0, 5.0)) 
            img_data = pix.tobytes("png")
            b64_img = base64.b64encode(img_data).decode('utf-8')
            images.append(f"data:image/png;base64,{b64_img}")
            
        return {"pages": images, "total_pages": len(doc)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8005)
