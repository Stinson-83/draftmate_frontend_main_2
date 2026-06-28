import logging
import os
from dotenv import load_dotenv

# Walk up and load the workspace root .env if it exists
_current_dir = os.path.dirname(os.path.abspath(__file__))
_env_path = os.path.abspath(os.path.join(_current_dir, "../..", ".env"))
if os.path.exists(_env_path):
    load_dotenv(_env_path)
else:
    load_dotenv()

ONLYOFFICE_API_URL = os.getenv("ONLYOFFICE_API_URL", "http://onlyoffice-server")

import logging
import hashlib
import json
from typing import Any, Dict, List, Optional, Set, Tuple
from urllib.parse import urlparse, urlunparse

import httpx
import jwt
import uvicorn
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
from dotenv import load_dotenv
from fastapi import FastAPI, Header, HTTPException, BackgroundTasks, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel, ConfigDict, Field, AliasChoices

import sys
import re
import time

current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.abspath(os.path.join(current_dir, "../"))
if parent_dir not in sys.path:
    sys.path.append(parent_dir)

deep_research_dir = os.path.abspath(os.path.join(current_dir, "../Deep_research"))
if deep_research_dir not in sys.path:
    sys.path.append(deep_research_dir)

try:
    from Deep_research.lex_bot.tools.pdf_processor import pdf_processor
    from Deep_research.lex_bot.tools.session_cache import get_session_cache
except ImportError:
    pdf_processor = None
    get_session_cache = None

try:
    from Deep_research.lex_bot.core.llm_factory import get_llm
except ImportError:
    get_llm = None

try:
    from Deep_research.lex_bot.tools import indian_kanoon_api as ik_api
except ImportError:
    ik_api = None

try:
    from Deep_research.lex_bot.tools.reranker import rerank_documents
except ImportError:
    rerank_documents = None

PLACEHOLDER_REGEX = re.compile(r'\b[A-Z][A-Z0-9_]{3,}\b')

def extract_and_cache_docx(file_path: str):
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        # Extract from paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text)
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            text_parts.append(p.text)
        full_text = "\n".join(text_parts)
        
        if full_text.strip() and pdf_processor and get_session_cache:
            chunks = pdf_processor.chunk_text(full_text)
            session_cache = get_session_cache()
            session_cache.set_file_chunks(file_path, chunks)
            logger.info(f"Background extraction complete for {file_path}. Chunks cached.")
    except Exception as e:
        logger.error(f"Background extraction failed: {e}")



from legal_draft import generate_legal_draft

logger = logging.getLogger(__name__)

app = FastAPI(title="DraftMate Drafter Service", version="2.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

AUTH_SERVICE_URL = os.getenv("AUTH_SERVICE_URL", "http://auth:8009")
JWT_SECRET = os.getenv("JWT_SECRET", "draftmate_jwt_production_signing_key_2026")
JWT_ALGORITHM = "HS256"


async def verify_token(authorization: Optional[str]) -> str:
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid Authorization header.")
    session_id = authorization.split(" ", 1)[1].strip()
    if not session_id:
        raise HTTPException(status_code=401, detail="Missing session token.")

    verify_url = f"{AUTH_SERVICE_URL.rstrip('/')}/verify_session/{session_id}"
    try:
        async with httpx.AsyncClient() as client:
            resp = await client.get(verify_url, timeout=10.0)
    except httpx.RequestError:
        if os.getenv("DEV_BYPASS_AUTH") == "true":
            logger.warning("DEV_BYPASS_AUTH enabled; bypassing auth service failure.")
            return "dev_counsel_bypass"
        raise HTTPException(status_code=503, detail="Auth service unavailable.")

    if resp.status_code != 200:
        raise HTTPException(status_code=401, detail="Invalid session.")

    try:
        data = resp.json()
    except Exception:
        raise HTTPException(status_code=502, detail="Auth service returned invalid JSON.")

    user_id = data.get("user_id")
    if not user_id:
        raise HTTPException(status_code=502, detail="Auth service response missing user_id.")
    return str(user_id)


def _jwt_encode(payload: Dict[str, Any]) -> str:
    token = jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)
    if isinstance(token, bytes):
        return token.decode("utf-8")
    return token


def _safe_basename_from_url(url: str) -> str:
    parsed = urlparse(url)
    name = os.path.basename(parsed.path or "")
    name = name.replace("\\", "/")
    name = os.path.basename(name)
    return name


def _normalize_download_url(download_source_url: str) -> str:
    parsed_source = urlparse(download_source_url)
    
    path = parsed_source.path
    if path.startswith("/onlyoffice/"):
        path = path.replace("/onlyoffice", "", 1)
        
    parsed_internal = urlparse(ONLYOFFICE_API_URL)
    
    normalized_url = urlunparse((
        parsed_internal.scheme or "http",
        parsed_internal.netloc or "onlyoffice-server",
        path,
        parsed_source.params,
        parsed_source.query,
        parsed_source.fragment
    ))
    return normalized_url


def _strip_json_block(text: str) -> str:
    if not text:
        return ""
    candidate = text.strip()
    if candidate.startswith("```"):
        candidate = re.sub(r"^```(?:json)?", "", candidate, flags=re.IGNORECASE).strip()
        if candidate.endswith("```"):
            candidate = candidate[:-3].strip()
    start = candidate.find("{")
    end = candidate.rfind("}")
    if start != -1 and end != -1 and end > start:
        return candidate[start : end + 1]
    return candidate


def _safe_json_loads(text: str) -> Optional[Dict[str, Any]]:
    if not text:
        return None
    try:
        payload = json.loads(_strip_json_block(text))
        if isinstance(payload, dict):
            return payload
    except Exception:
        return None
    return None


def _llm_text_response(llm: Any, prompt: str) -> str:
    if llm is None:
        return ""
    try:
        if hasattr(llm, "invoke"):
            result = llm.invoke(prompt)
        elif hasattr(llm, "predict"):
            result = llm.predict(prompt)
        else:
            result = llm(prompt)
    except Exception as exc:
        logger.warning(f"LLM invocation failed: {exc}")
        return ""

    if isinstance(result, str):
        return result
    if hasattr(result, "content"):
        return str(result.content)
    return str(result)


def _normalize_answer_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, dict):
        parts = []
        for key in sorted(value.keys()):
            normalized = _normalize_answer_text(value[key])
            if normalized:
                parts.append(f"{key}: {normalized}")
        return " | ".join(parts).strip()
    if isinstance(value, list):
        return ", ".join([_normalize_answer_text(v) for v in value if _normalize_answer_text(v)])
    return str(value).strip()


def _flatten_answers(accumulated_answers: Dict[str, Any]) -> str:
    if not accumulated_answers:
        return ""
    pieces = []
    for key in sorted(accumulated_answers.keys()):
        value = _normalize_answer_text(accumulated_answers.get(key))
        if value:
            pieces.append(f"{key}: {value}")
    return "\n".join(pieces)


def _build_clarifying_question(question_id: str, prompt: str, options: List[Tuple[str, str]]) -> ClarifyingQuestion:
    return ClarifyingQuestion(
        id=question_id,
        question=prompt,
        type="single",
        options=[ClarifyingOption(label=label, value=value) for label, value in options],
        required=True,
    )


def _heuristic_intake_analysis(initial_prompt: str, accumulated_answers: Dict[str, Any], current_round_index: int) -> Dict[str, Any]:
    text = f"{initial_prompt}\n{_flatten_answers(accumulated_answers)}".lower()

    # Look for explicit answers in accumulated_answers
    ans_doc_type = None
    ans_jurisdiction = None
    ans_position = None

    for k, v in accumulated_answers.items():
        k_lower = str(k).lower()
        v_str = str(v).strip()
        if not v_str:
            continue
        if "document" in k_lower or "preparing" in k_lower or "doc_type" in k_lower:
            ans_doc_type = v_str
        if "jurisdiction" in k_lower or "govern" in k_lower:
            ans_jurisdiction = v_str
        if "representing" in k_lower or "side" in k_lower or "position" in k_lower:
            ans_position = v_str

    document_type = "Legal Document"
    if any(token in text for token in ["lease", "tenancy"]):
        document_type = "Lease / Tenancy Agreement"
    elif any(token in text for token in ["employment", "employee", "employer"]):
        document_type = "Employment Agreement"
    elif any(token in text for token in ["share", "equity", "shareholders"]):
        document_type = "Shareholders Agreement"
    elif any(token in text for token in ["service", "consulting", "vendor"]):
        document_type = "Services Agreement"
    elif any(token in text for token in ["nda", "non-disclosure", "confidential"]):
        document_type = "NDA / Confidentiality Agreement"
    elif any(token in text for token in ["petition", "writ", "appeal", "suit"]):
        document_type = "Litigation Filing"

    if ans_doc_type:
        document_type = ans_doc_type

    jurisdiction = "Not specified"
    for candidate in ["india", "indian", "delhi", "mumbai", "maharashtra", "karnataka", "tamil nadu", "california", "new york"]:
        if candidate in text:
            jurisdiction = candidate.title()
            break

    if ans_jurisdiction:
        jurisdiction = ans_jurisdiction

    rep_position = "Not specified"
    for candidate in [
        "for the claimant",
        "for the plaintiff",
        "for the petitioner",
        "for the respondent",
        "for the defendant",
        "for the seller",
        "for the buyer",
        "for the licensor",
        "for the licensee",
    ]:
        if candidate in text:
            rep_position = candidate.replace("for the ", "").title()
            break

    if ans_position:
        rep_position = ans_position

    questions: List[ClarifyingQuestion] = []
    missing_key_inputs = []
    if jurisdiction == "Not specified" and not ans_jurisdiction:
        missing_key_inputs.append("jurisdiction")
    if document_type == "Legal Document" and not ans_doc_type:
        missing_key_inputs.append("document_type")
    if rep_position == "Not specified" and not ans_position:
        missing_key_inputs.append("representation_position")
    if not accumulated_answers:
        missing_key_inputs.append("commercial_terms")

    if missing_key_inputs:
        if "jurisdiction" in missing_key_inputs:
            questions.append(
                _build_clarifying_question(
                    "jurisdiction",
                    "Which jurisdiction should govern this draft?",
                    [("India", "India"), ("United States", "United States"), ("Other / mixed", "Other / mixed")],
                )
            )
        if "document_type" in missing_key_inputs:
            questions.append(
                _build_clarifying_question(
                    "document_type",
                    "What kind of document are we preparing?",
                    [
                        (document_type, document_type),
                        ("Contract / Agreement", "Contract / Agreement"),
                        ("Litigation / Court filing", "Litigation / Court filing"),
                        ("Other", "Other"),
                    ],
                )
            )
        if "representation_position" in missing_key_inputs:
            questions.append(
                _build_clarifying_question(
                    "position",
                    "Which side are we representing?",
                    [
                        ("Party A / initiator", "Party A / initiator"),
                        ("Party B / responding party", "Party B / responding party"),
                        ("Buyer / recipient", "Buyer / recipient"),
                        ("Seller / provider", "Seller / provider"),
                    ],
                )
            )

    basis = DraftBasis(
        document_type=document_type,
        jurisdiction=jurisdiction,
        representation_position=rep_position,
        key_legal_positions=[
            "Use market-standard allocation of risk where the prompt is silent.",
            "Preserve internal consistency across definitions, remedies, and termination rights.",
        ],
    )
    assumptions = [
        "Market-standard boilerplate will be used for missing administrative clauses.",
        "Ambiguous commercial inputs will be resolved conservatively in favor of internal consistency.",
    ]

    sufficient = len(questions) == 0
    return {
        "sufficiency_met": sufficient,
        "questions": questions[:5],
        "draft_summary": DraftSummary(basis=basis, assumptions=assumptions) if sufficient else None,
        "validation_metadata": {
            "document_type": document_type,
            "jurisdiction": jurisdiction,
            "representation_position": rep_position,
            "current_round_index": current_round_index,
            "detected_missing_inputs": missing_key_inputs,
            "source": "heuristic_fallback",
        },
    }


def _build_intake_prompt(initial_prompt: str, accumulated_answers: Dict[str, Any], current_round_index: int) -> str:
    answers_text = _flatten_answers(accumulated_answers)
    return f"""
You are a Senior Transactional Lawyer performing intake analysis for a legal drafting workflow.

Follow these steps:
1. Document Identification.
2. Requirement Extraction.
3. Gap Analysis.

If critical structural or commercial risk information is missing, return up to 5 concise clarifying questions.
Prefer multiple-choice questions with 3-4 options.
If the matter is sufficiently clear or market-standard provisions can close the gaps, return sufficiency_met true,
and include:
- Step 6 assumptions
- Step 8 draft summary basis

Return JSON only with this schema:
{{
  "sufficiency_met": boolean,
  "questions": [{{"id": string, "question": string, "type": "single"|"multiple", "options": [{{"label": string, "value": string}}]}}],
  "draft_summary": {{
    "basis": {{
      "document_type": string,
      "jurisdiction": string,
      "representation_position": string,
      "key_legal_positions": [string]
    }},
    "assumptions": [string]
  }},
  "validation_metadata": object
}}

Current round index: {current_round_index}
Initial prompt:
{initial_prompt}

Accumulated answers:
{answers_text or "[none]"}
""".strip()


def _run_intake_orchestrator(initial_prompt: str, accumulated_answers: Dict[str, Any], current_round_index: int) -> IntakeAnalyzeResponse:
    if get_llm is not None:
        try:
            llm = get_llm(mode="reasoning")
            prompt = _build_intake_prompt(initial_prompt, accumulated_answers, current_round_index)
            raw = _llm_text_response(llm, prompt)
            parsed = _safe_json_loads(raw) or {}
            sufficiency_met = bool(parsed.get("sufficiency_met"))
            questions_payload = parsed.get("questions") or []
            questions: List[ClarifyingQuestion] = []
            for idx, item in enumerate(questions_payload[:5]):
                if not isinstance(item, dict):
                    continue
                options = item.get("options") or []
                questions.append(
                    ClarifyingQuestion(
                        id=str(item.get("id") or f"q_{current_round_index}_{idx}"),
                        question=str(item.get("question") or ""),
                        type=str(item.get("type") or "single"),
                        options=[
                            ClarifyingOption(
                                label=str(opt.get("label") or opt.get("value") or opt),
                                value=str(opt.get("value") or opt.get("label") or opt),
                            )
                            for opt in options
                            if isinstance(opt, (dict, str))
                        ],
                        required=True,
                    )
                )

            draft_summary_payload = parsed.get("draft_summary") or {}
            draft_summary = None
            if sufficiency_met and isinstance(draft_summary_payload, dict):
                basis_payload = draft_summary_payload.get("basis") or {}
                basis = DraftBasis(
                    document_type=str(basis_payload.get("document_type") or "Legal Document"),
                    jurisdiction=str(basis_payload.get("jurisdiction") or "Not specified"),
                    representation_position=str(basis_payload.get("representation_position") or "Not specified"),
                    key_legal_positions=[str(x) for x in (basis_payload.get("key_legal_positions") or []) if str(x).strip()],
                )
                draft_summary = DraftSummary(
                    basis=basis,
                    assumptions=[str(x) for x in (draft_summary_payload.get("assumptions") or []) if str(x).strip()],
                )

            if sufficiency_met and draft_summary is None:
                fallback = _heuristic_intake_analysis(initial_prompt, accumulated_answers, current_round_index)
                draft_summary = fallback["draft_summary"]
                parsed_metadata = fallback["validation_metadata"]
            else:
                parsed_metadata = parsed.get("validation_metadata") if isinstance(parsed.get("validation_metadata"), dict) else {}

            if not parsed_metadata:
                parsed_metadata = {}
            parsed_metadata.setdefault("source", "llm")
            parsed_metadata.setdefault("current_round_index", current_round_index)
            return IntakeAnalyzeResponse(
                sufficiency_met=sufficiency_met or (draft_summary is not None and not questions),
                current_round_index=current_round_index,
                next_round_index=current_round_index + 1,
                questions=questions[:5],
                draft_summary=draft_summary,
                validation_metadata=parsed_metadata,
            )
        except Exception as exc:
            logger.warning(f"LLM intake orchestrator failed, using heuristic fallback: {exc}")

    fallback = _heuristic_intake_analysis(initial_prompt, accumulated_answers, current_round_index)
    return IntakeAnalyzeResponse(
        sufficiency_met=bool(fallback["sufficiency_met"]),
        current_round_index=current_round_index,
        next_round_index=current_round_index + 1,
        questions=fallback["questions"][:5],
        draft_summary=fallback["draft_summary"],
        validation_metadata=fallback["validation_metadata"],
    )


def _extract_case_tokens(raw_text: str, document_vehicle: Optional[str], jurisdiction_context: Optional[str]) -> Dict[str, str]:
    base_text = (raw_text or "").strip()
    if get_llm is not None:
        try:
            llm = get_llm(mode="fast")
            prompt = f"""
Extract the following from the highlighted legal text:
- Legal Issue
- Jurisdiction Context
- Document Vehicle
- Core Proposition

Return JSON only:
{{
  "legal_issue": string,
  "jurisdiction_context": string,
  "document_vehicle": string,
  "core_proposition": string
}}

Highlighted text:
{base_text}

Known context:
document_vehicle: {document_vehicle or ""}
jurisdiction_context: {jurisdiction_context or ""}
""".strip()
            raw = _llm_text_response(llm, prompt)
            parsed = _safe_json_loads(raw) or {}
            return {
                "legal_issue": str(parsed.get("legal_issue") or base_text[:220] or "legal issue"),
                "jurisdiction_context": str(parsed.get("jurisdiction_context") or jurisdiction_context or "unspecified"),
                "document_vehicle": str(parsed.get("document_vehicle") or document_vehicle or "unspecified"),
                "core_proposition": str(parsed.get("core_proposition") or base_text[:320] or "unspecified"),
            }
        except Exception as exc:
            logger.warning(f"Case token extraction LLM failed, using heuristic fallback: {exc}")

    lowered = base_text.lower()
    if any(word in lowered for word in ["appeal", "supreme court", "high court", "trial court", "writ"]):
        issue = "Judicial precedent and procedural posture"
    elif any(word in lowered for word in ["lease", "rent", "tenant", "landlord"]):
        issue = "Property or tenancy dispute"
    elif any(word in lowered for word in ["employment", "dismissal", "termination"]):
        issue = "Employment dispute"
    elif any(word in lowered for word in ["contract", "agreement", "breach", "indemnity"]):
        issue = "Contract interpretation or breach"
    else:
        issue = base_text[:160] or "legal issue"

    return {
        "legal_issue": issue,
        "jurisdiction_context": jurisdiction_context or "unspecified",
        "document_vehicle": document_vehicle or "unspecified",
        "core_proposition": base_text[:320] or "unspecified",
    }


def _case_search_queries(tokens: Dict[str, str]) -> List[str]:
    queries = [
        " ".join(part for part in [tokens.get("legal_issue"), tokens.get("jurisdiction_context")] if part and part != "unspecified").strip(),
        " ".join(part for part in [tokens.get("core_proposition"), tokens.get("jurisdiction_context")] if part and part != "unspecified").strip(),
        " ".join(part for part in [tokens.get("document_vehicle"), tokens.get("legal_issue")] if part and part != "unspecified").strip(),
    ]
    seen: Set[str] = set()
    final_queries: List[str] = []
    for query in queries:
        normalized = query.strip()
        if normalized and normalized.lower() not in seen:
            seen.add(normalized.lower())
            final_queries.append(normalized)
    if tokens.get("legal_issue") and tokens["legal_issue"].lower() not in seen:
        final_queries.append(tokens["legal_issue"])
    return final_queries[:5]


def _court_label(candidate: Dict[str, Any]) -> str:
    for key in ("court", "docsource", "source", "doctype"):
        value = candidate.get(key)
        if isinstance(value, str) and value.strip():
            return value.strip()
    return "Court not specified"


def _case_weight(candidate: Dict[str, Any], tokens: Dict[str, str]) -> float:
    haystack = " ".join(
        str(candidate.get(key, "")) for key in ("title", "snippet", "citation", "court", "docsource", "source")
    ).lower()
    weight = 0.0
    for needle in [
        tokens.get("legal_issue", ""),
        tokens.get("core_proposition", ""),
        tokens.get("jurisdiction_context", ""),
        tokens.get("document_vehicle", ""),
    ]:
        needle = (needle or "").lower().strip()
        if needle and needle != "unspecified" and needle in haystack:
            weight += 2.0
    if "supreme" in haystack or "high court" in haystack:
        weight += 1.25
    if candidate.get("citation"):
        weight += 0.5
    return weight


def _summarize_case_candidate(candidate: Dict[str, Any], tokens: Dict[str, str]) -> CaseSearchItem:
    title = str(candidate.get("title") or candidate.get("case_name") or "Unknown case")
    citation = str(candidate.get("citation") or candidate.get("docid") or "")
    snippet = str(candidate.get("snippet") or candidate.get("summary") or "")
    court = _court_label(candidate)
    relevance = candidate.get("relevance_justification")
    if not isinstance(relevance, str) or not relevance.strip():
        relevance = (
            f"Matched the issue profile around {tokens.get('legal_issue', 'the highlighted issue')} "
            f"and aligns with the stated proposition."
        )
    holding = candidate.get("holding_summary")
    if not isinstance(holding, str) or not holding.strip():
        holding = snippet[:320] if snippet else "Holding summary unavailable from the retrieved source."
    return CaseSearchItem(
        case_name=title,
        court=court,
        citation=citation,
        relevance_justification=relevance.strip(),
        holding_summary=holding.strip(),
        source_url=str(candidate.get("url") or candidate.get("source_url") or "") or None,
        docid=str(candidate.get("docid") or "") or None,
    )


def _fetch_case_candidates(query: str, limit: int) -> List[Dict[str, Any]]:
    if ik_api is None:
        return []
    try:
        results = ik_api.search(query, max_results=max(5, min(limit, 10)))
    except Exception as exc:
        logger.warning(f"Indian Kanoon search failed for query '{query}': {exc}")
        return []
    if isinstance(results, list):
        return [item for item in results if isinstance(item, dict)]
    return []


def _rank_case_candidates(candidates: List[Dict[str, Any]], tokens: Dict[str, str]) -> List[Dict[str, Any]]:
    if not candidates:
        return []

    if rerank_documents is not None:
        try:
            query = " ".join([tokens.get("legal_issue", ""), tokens.get("core_proposition", ""), tokens.get("jurisdiction_context", "")]).strip()
            docs = []
            for candidate in candidates:
                docs.append(
                    {
                        "text": " ".join(str(candidate.get(key, "")) for key in ("title", "snippet", "citation", "docsource")).strip(),
                        "metadata": candidate,
                    }
                )
            ranked = rerank_documents(query=query, candidates=docs, top_n=min(10, len(docs)))
            if isinstance(ranked, list) and ranked:
                ranked_candidates: List[Dict[str, Any]] = []
                for item in ranked:
                    if isinstance(item, dict):
                        metadata = item.get("metadata") if isinstance(item.get("metadata"), dict) else None
                        ranked_candidates.append(metadata or item)
                if ranked_candidates:
                    candidates = ranked_candidates
        except Exception as exc:
            logger.warning(f"Reranking failed, using heuristic ranking: {exc}")

    unique: Dict[str, Dict[str, Any]] = {}
    for candidate in candidates:
        docid = str(candidate.get("docid") or candidate.get("id") or candidate.get("citation") or candidate.get("title") or "").strip()
        if not docid:
            continue
        score = _case_weight(candidate, tokens)
        if docid not in unique or score > unique[docid].get("_score", -1):
            candidate["_score"] = score
            unique[docid] = candidate
    return sorted(unique.values(), key=lambda item: item.get("_score", 0.0), reverse=True)

def create_content_control_run(paragraph, placeholder_tag: str, default_text: str):
    sdt = OxmlElement("w:sdt")


    sdt_pr = OxmlElement("w:sdtPr")

    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), placeholder_tag)
    sdt_pr.append(tag)

    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), placeholder_tag)
    sdt_pr.append(alias)

    showing = OxmlElement("w:showingPlcHdr")
    sdt_pr.append(showing)

    sdt.append(sdt_pr)

    sdt_content = OxmlElement("w:sdtContent")

    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")
    r_pr.append(r_fonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    r_pr.append(sz)

    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), "24")
    r_pr.append(sz_cs)

    r.append(r_pr)

    t = OxmlElement("w:t")
    t.text = f"[{default_text}]"
    r.append(t)

    sdt_content.append(r)
    sdt.append(sdt_content)

    paragraph._p.append(sdt)


def _token_core_segments(token: str) -> Tuple[str, str, str]:
    if not token:
        return "", "", ""
    start = 0
    while start < len(token) and not (token[start].isalnum() or token[start] == "_"):
        start += 1
    end = len(token)
    while end > start and not (token[end - 1].isalnum() or token[end - 1] == "_"):
        end -= 1
    return token[:start], token[start:end], token[end:]


def _apply_run_style(run, font_size: Pt, bold: bool):
    run.font.name = "Times New Roman"
    run.font.size = font_size
    run.bold = bold


def build_docx_with_controls(ai_data: dict, file_target_name: str) -> str:
    shared_storage_path = os.getenv("SHARED_STORAGE_PATH")
    if not shared_storage_path:
        raise ValueError("SHARED_STORAGE_PATH is not set.")

    os.makedirs(shared_storage_path, exist_ok=True)

    safe_name = (file_target_name or "").strip() or "draftmate_draft.docx"
    if not safe_name.lower().endswith(".docx"):
        safe_name = safe_name + ".docx"
    output_path = os.path.join(shared_storage_path, safe_name)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    metadata = ai_data.get("metadata") or {}
    placeholders_list = metadata.get("placeholders_detected") or []
    placeholders: Set[str] = set(x for x in placeholders_list if isinstance(x, str))

    content_blocks = ai_data.get("content") or []
    for block in content_blocks:
        if not isinstance(block, dict):
            continue

        element_type = block.get("element_type")
        text = block.get("text") or ""
        if not isinstance(text, str):
            text = str(text)

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.5

        if element_type == "header_block":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            font_size = Pt(14)
            bold = True
        elif element_type == "paragraph":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            font_size = Pt(12)
            bold = False
        elif element_type == "heading_1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            font_size = Pt(13)
            bold = True
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            font_size = Pt(12)
            bold = False

        tokens = text.split(" ")
        for idx, token in enumerate(tokens):
            if idx > 0:
                space_run = paragraph.add_run(" ")
                _apply_run_style(space_run, font_size=font_size, bold=bold)

            prefix, core, suffix = _token_core_segments(token)
            if core and core in placeholders:
                if prefix:
                    r = paragraph.add_run(prefix)
                    _apply_run_style(r, font_size=font_size, bold=bold)

                create_content_control_run(paragraph, placeholder_tag=core, default_text=core)

                if suffix:
                    r = paragraph.add_run(suffix)
                    _apply_run_style(r, font_size=font_size, bold=bold)
            else:
                r = paragraph.add_run(token)
                _apply_run_style(r, font_size=font_size, bold=bold)

    doc.save(output_path)
    return output_path


class DraftCompileRequest(BaseModel):
    case_context: Optional[str] = None
    case_metadata_context: Optional[List[Dict[str, Any]]] = None
    legal_documents: Optional[str] = None
    document_type: str = "Legal Document"
    file_target_name: str = "draftmate_draft.docx"


class ForceSaveRequest(BaseModel):
    document_key: str


class IntakeAnalyzeRequest(BaseModel):
    model_config = ConfigDict(populate_by_name=True, extra="allow")

    initial_prompt: str = Field(
        default="",
        validation_alias=AliasChoices("initial_prompt", "prompt", "case_context", "user_prompt"),
    )
    current_round_index: int = 0
    accumulated_answers: Dict[str, Any] = Field(
        default_factory=dict,
        validation_alias=AliasChoices("accumulated_answers", "answers"),
    )


class ClarifyingOption(BaseModel):
    model_config = ConfigDict(extra="forbid")

    label: str
    value: str


class ClarifyingQuestion(BaseModel):
    model_config = ConfigDict(extra="forbid")

    id: str
    question: str
    type: str = "single"
    options: List[ClarifyingOption] = Field(default_factory=list)
    required: bool = True


class DraftBasis(BaseModel):
    model_config = ConfigDict(extra="forbid")

    document_type: str
    jurisdiction: str
    representation_position: str
    key_legal_positions: List[str] = Field(default_factory=list)


class DraftSummary(BaseModel):
    model_config = ConfigDict(extra="forbid")

    basis: DraftBasis
    assumptions: List[str] = Field(default_factory=list)


class IntakeAnalyzeResponse(BaseModel):
    model_config = ConfigDict(extra="forbid")

    sufficiency_met: bool
    current_round_index: int
    next_round_index: int
    questions: List[ClarifyingQuestion] = Field(default_factory=list)
    draft_summary: Optional[DraftSummary] = None
    validation_metadata: Dict[str, Any] = Field(default_factory=dict)


class CaseSearchRequest(BaseModel):
    model_config = ConfigDict(populate_by_name=True, extra="allow")

    raw_text: str = Field(
        default="",
        validation_alias=AliasChoices("raw_text", "highlighted_text", "selection", "query", "text"),
    )
    document_vehicle: Optional[str] = None
    jurisdiction_context: Optional[str] = None
    limit: int = 8


class CaseSearchItem(BaseModel):
    model_config = ConfigDict(extra="forbid")

    case_name: str
    court: str
    citation: str
    relevance_justification: str
    holding_summary: str
    source_url: Optional[str] = None
    docid: Optional[str] = None


class CaseSearchResponse(BaseModel):
    model_config = ConfigDict(extra="forbid")

    tokens: Dict[str, str]
    cases: List[CaseSearchItem] = Field(default_factory=list)
    source: str = "indian_kanoon"


@app.post("/v2/draft/intake/analyze", response_model=IntakeAnalyzeResponse)
async def intake_analyze(request: IntakeAnalyzeRequest, authorization: Optional[str] = Header(default=None)):
    await verify_token(authorization)
    return _run_intake_orchestrator(
        initial_prompt=request.initial_prompt,
        accumulated_answers=request.accumulated_answers,
        current_round_index=request.current_round_index,
    )


@app.post("/v2/research/cases", response_model=CaseSearchResponse)
async def research_cases(request: CaseSearchRequest, authorization: Optional[str] = Header(default=None)):
    await verify_token(authorization)

    tokens = _extract_case_tokens(
        raw_text=request.raw_text,
        document_vehicle=request.document_vehicle,
        jurisdiction_context=request.jurisdiction_context,
    )
    queries = _case_search_queries(tokens)

    candidates: List[Dict[str, Any]] = []
    for query in queries:
        candidates.extend(_fetch_case_candidates(query, request.limit))

    ranked_candidates = _rank_case_candidates(candidates, tokens)
    top_candidates = ranked_candidates[: max(5, min(request.limit, 10))]
    cases = [_summarize_case_candidate(candidate, tokens) for candidate in top_candidates]

    return CaseSearchResponse(tokens=tokens, cases=cases)


@app.get("/")
def root():
    return {"service": "drafter-service", "status": "ok"}


class VariableSyncRequest(BaseModel):
    filename: str
    placeholder_tag: str
    new_value: str


@app.post("/v2/draft/variable/sync")
async def sync_variable_value(request: VariableSyncRequest, authorization: Optional[str] = Header(default=None)):
    await verify_token(authorization)
    
    shared_storage_path = os.getenv("SHARED_STORAGE_PATH")
    if not shared_storage_path:
        raise HTTPException(status_code=500, detail="SHARED_STORAGE_PATH is not set.")
        
    file_path = os.path.join(shared_storage_path, os.path.basename(request.filename))
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Target document not found.")
        
    try:
        from docx import Document
        doc = Document(file_path)
        
        modified = False
        body_elements = doc.element.body
        for sdt in body_elements.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdt'):
            tag_elem = sdt.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tag')
            if tag_elem is not None and tag_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') == request.placeholder_tag:
                t_elements = sdt.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                if t_elements:
                    t_elements[0].text = request.new_value
                    for extra_t in t_elements[1:]:
                        extra_t.text = ""
                    modified = True
                    
        if modified:
            doc.save(file_path)
            return {"status": "synchronized", "updated": True}
            
        return {"status": "unchanged", "updated": False}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"XML modification failed: {str(e)}")


@app.get("/v2/draft/serve/{filename}")
def serve_draft_file(filename: str):
    shared_storage_path = os.getenv("SHARED_STORAGE_PATH")
    if not shared_storage_path:
        raise HTTPException(status_code=500, detail="SHARED_STORAGE_PATH is not set.")

    safe_name = os.path.basename((filename or "").replace("\\", "/"))
    if not safe_name:
        raise HTTPException(status_code=400, detail="Invalid filename.")

    file_path = os.path.join(shared_storage_path, safe_name)
    if not os.path.isfile(file_path):
        raise HTTPException(status_code=404, detail="File not found.")

    return FileResponse(
        path=file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=safe_name,
    )


@app.get("/v2/draft/serve/{draft_id}/{filename}")
async def serve_draft(draft_id: str, filename: str):
    shared_storage_path = os.getenv("SHARED_STORAGE_PATH")
    if not shared_storage_path:
        raise HTTPException(status_code=500, detail="SHARED_STORAGE_PATH is not set.")

    safe_draft_id = os.path.basename(draft_id.replace("\\", "/"))
    safe_name = os.path.basename((filename or "").replace("\\", "/"))
    if not safe_name or not safe_draft_id:
        raise HTTPException(status_code=400, detail="Invalid path parameter.")

    file_path = os.path.join(shared_storage_path, safe_draft_id, safe_name)
    if not os.path.isfile(file_path):
        root_path = os.path.join(shared_storage_path, safe_name)
        if os.path.isfile(root_path):
            file_path = root_path
        else:
            raise HTTPException(status_code=404, detail="File not found.")

    return FileResponse(
        path=file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=safe_name,
    )


@app.post("/v2/draft/compile")
async def compile_draft(request: DraftCompileRequest, authorization: Optional[str] = Header(default=None)):
    try:
        user_id = await verify_token(authorization)

        context_parts: List[str] = []
        if request.case_metadata_context:
            for item in request.case_metadata_context:
                if isinstance(item, dict):
                    for k, v in item.items():
                        context_parts.append(f"{k}: {v}")
        if request.case_context:
            context_parts.append(request.case_context)
        case_context = "\n".join([p for p in context_parts if p]).strip()
        if not case_context:
            raise ValueError("case_context is required.")

        ai_data = generate_legal_draft(
            case_context=case_context,
            legal_documents=request.legal_documents,
            document_type=request.document_type,
        )
        
        shared_storage_path = os.getenv("SHARED_STORAGE_PATH")
        if not shared_storage_path:
            raise HTTPException(status_code=500, detail="SHARED_STORAGE_PATH is not set.")

        import uuid
        draft_id = str(uuid.uuid4())
        draft_dir = os.path.join(shared_storage_path, draft_id)
        os.makedirs(draft_dir, exist_ok=True)

        output_path = build_docx_with_controls(ai_data=ai_data, file_target_name=request.file_target_name)
        file_name = os.path.basename(output_path)
        
        # Move generated file to the sandboxed path
        sandboxed_path = os.path.join(draft_dir, file_name)
        os.rename(output_path, sandboxed_path)

        document_key = hashlib.sha256(draft_id.encode("utf-8")).hexdigest()

        # Copy file to lex_bot upload directory
        current_dir = os.path.dirname(os.path.abspath(__file__))
        lex_bot_upload_dir = os.path.abspath(os.path.join(current_dir, "../Deep_research/lex_bot/data/uploads"))
        os.makedirs(lex_bot_upload_dir, exist_ok=True)
        lex_bot_path = os.path.join(lex_bot_upload_dir, file_name)
        with open(sandboxed_path, "rb") as sf:
            with open(lex_bot_path, "wb") as lf:
                lf.write(sf.read())

        metadata = ai_data.get("metadata") or {}
        placeholders_list = metadata.get("placeholders_detected") or []

        # Register draft metadata in PostgreSQL db via auth service
        try:
            async with httpx.AsyncClient() as client:
                await client.post(
                    f"{AUTH_SERVICE_URL.rstrip('/')}/internal/draft/register",
                    json={
                        "draft_id": draft_id,
                        "name": request.file_target_name or file_name,
                        "filename": file_name,
                        "document_key": document_key,
                        "created_by": user_id,
                        "variables_detected": placeholders_list,
                        "status": "In progress"
                    },
                    timeout=10.0
                )
        except Exception as reg_err:
            logger.warning(f"Failed to register draft in DB: {reg_err}")

        params: Dict[str, Any] = {
            "document": {
                "fileType": "docx",
                "key": document_key,
                "title": file_name,
                "url": f"http://drafter-service:8003/v2/draft/serve/{draft_id}/{file_name}",
                "permissions": {"edit": True, "download": True, "print": True},
            },
            "documentType": "word",
            "editorConfig": {
                "callbackUrl": f"http://drafter-service:8003/v2/draft/callback/{draft_id}",
                "mode": "edit",
                "user": {
                    "id": user_id,
                    "name": f"User {user_id[:4]}" if len(user_id) >= 4 else f"User {user_id}"
                },
                "coauthoring": {
                    "mode": "fast",
                    "change": True
                },
                "customization": {
                    "forcesave": True,
                    "chat": True,
                    "uiTheme": "theme-light",
                    "logo": {
                        "image": "",
                        "imageDark": "",
                        "url": ""
                    }
                },
                "plugins": {
                    "autostart": [
                        "asc.{43d1a84f-e274-4b53-a55e-3363f8db1f34}"
                    ],
                    "pluginsData": []
                }
            },
        }
        params["token"] = _jwt_encode(params)
        params["documentKey"] = document_key
        params["filename"] = file_name
        params["variablesDetected"] = placeholders_list
        params["draftId"] = draft_id

        return params
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.exception("Draft compilation failed.")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/v2/draft/create")
async def create_empty_draft(authorization: Optional[str] = Header(default=None)):
    try:
        user_id = await verify_token(authorization)

        shared_storage_path = os.getenv("SHARED_STORAGE_PATH")
        if not shared_storage_path:
            raise HTTPException(status_code=500, detail="SHARED_STORAGE_PATH is not set.")

        import uuid
        draft_id = str(uuid.uuid4())
        draft_dir = os.path.join(shared_storage_path, draft_id)
        os.makedirs(draft_dir, exist_ok=True)

        file_name = f"Empty_Draft_{int(time.time())}.docx"
        output_path = os.path.join(draft_dir, file_name)

        from docx import Document
        doc = Document()
        doc.add_paragraph("")
        doc.save(output_path)

        with open(output_path, "rb") as f:
            file_bytes = f.read()

        document_key = hashlib.sha256(draft_id.encode("utf-8")).hexdigest()

        # Copy empty file to lex_bot upload directory
        current_dir = os.path.dirname(os.path.abspath(__file__))
        lex_bot_upload_dir = os.path.abspath(os.path.join(current_dir, "../Deep_research/lex_bot/data/uploads"))
        os.makedirs(lex_bot_upload_dir, exist_ok=True)
        lex_bot_path = os.path.join(lex_bot_upload_dir, file_name)
        with open(lex_bot_path, "wb") as f:
            f.write(file_bytes)

        # Register draft metadata in PostgreSQL db via auth service
        try:
            async with httpx.AsyncClient() as client:
                await client.post(
                    f"{AUTH_SERVICE_URL.rstrip('/')}/internal/draft/register",
                    json={
                        "draft_id": draft_id,
                        "name": file_name,
                        "filename": file_name,
                        "document_key": document_key,
                        "created_by": user_id,
                        "variables_detected": [],
                        "status": "In progress"
                    },
                    timeout=10.0
                )
        except Exception as reg_err:
            logger.warning(f"Failed to register draft in DB: {reg_err}")

        params: Dict[str, Any] = {
            "document": {
                "fileType": "docx",
                "key": document_key,
                "title": file_name,
                "url": f"http://drafter-service:8003/v2/draft/serve/{draft_id}/{file_name}",
                "permissions": {"edit": True, "download": True, "print": True},
            },
            "documentType": "word",
            "editorConfig": {
                "callbackUrl": f"http://drafter-service:8003/v2/draft/callback/{draft_id}",
                "mode": "edit",
                "user": {
                    "id": user_id,
                    "name": f"User {user_id[:4]}" if len(user_id) >= 4 else f"User {user_id}"
                },
                "coauthoring": {
                    "mode": "fast",
                    "change": True
                },
                "customization": {
                    "forcesave": True,
                    "chat": True,
                    "uiTheme": "theme-light",
                    "logo": {
                        "image": "",
                        "imageDark": "",
                        "url": ""
                    }
                },
            },
        }
        params["token"] = _jwt_encode(params)
        params["documentKey"] = document_key
        params["filename"] = file_name
        params["variablesDetected"] = []
        params["draftId"] = draft_id

        return params
    except Exception as e:
        logger.exception("Failed to create empty draft.")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/v2/draft/upload")
async def upload_draft(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    session_id: Optional[str] = Form(None),
    authorization: Optional[str] = Header(default=None)
):
    try:
        user_id = await verify_token(authorization)
        
        shared_storage_path = os.getenv("SHARED_STORAGE_PATH")
        if not shared_storage_path:
            raise HTTPException(status_code=500, detail="SHARED_STORAGE_PATH is not set.")
            
        import uuid
        draft_id = str(uuid.uuid4())
        draft_dir = os.path.join(shared_storage_path, draft_id)
        os.makedirs(draft_dir, exist_ok=True)
        
        file_name = file.filename or "uploaded_draft.docx"
        safe_name = file_name.strip()
        
        file_bytes = await file.read()
        is_pdf = safe_name.lower().endswith(".pdf") or file_bytes.startswith(b"%PDF")
        
        if is_pdf:
            if safe_name.lower().endswith(".pdf"):
                safe_name = safe_name[:-4] + ".docx"
            else:
                safe_name = safe_name + ".docx"
        else:
            if not safe_name.lower().endswith(".docx"):
                if "." in safe_name:
                    parts = safe_name.rsplit(".", 1)
                    safe_name = f"{parts[0]}_{int(time.time())}.docx"
                else:
                    safe_name = f"{safe_name}_{int(time.time())}.docx"
                
        output_path = os.path.join(draft_dir, safe_name)
        
        if is_pdf:
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                temp_pdf.write(file_bytes)
                temp_pdf_path = temp_pdf.name
            
            try:
                from pdf2docx import Converter
                cv = Converter(temp_pdf_path)
                cv.convert(output_path, start=0, end=None)
                cv.close()
            except Exception as conv_e:
                logger.error(f"pdf2docx conversion failed: {conv_e}")
                # Fallback: simple text extraction via fitz (PyMuPDF)
                try:
                    import fitz
                    from docx import Document
                    pdf_doc = fitz.open(temp_pdf_path)
                    doc = Document()
                    for page_num in range(len(pdf_doc)):
                        page = pdf_doc[page_num]
                        text = page.get_text()
                        if page_num > 0:
                            doc.add_page_break()
                        if text.strip():
                            for para_text in text.split('\n\n'):
                                if para_text.strip():
                                    doc.add_paragraph(para_text.strip())
                    doc.save(output_path)
                except Exception as fitz_e:
                    logger.error(f"Fallback fitz conversion also failed: {fitz_e}")
                    raise HTTPException(status_code=500, detail="Failed to convert PDF to DOCX.")
            finally:
                try:
                    os.unlink(temp_pdf_path)
                except Exception:
                    pass
        else:
            with open(output_path, "wb") as f:
                f.write(file_bytes)
            
        # Read the DOCX bytes
        with open(output_path, "rb") as f:
            docx_bytes = f.read()

        # Copy file to lex_bot upload directory
        current_dir = os.path.dirname(os.path.abspath(__file__))
        lex_bot_upload_dir = os.path.abspath(os.path.join(current_dir, "../Deep_research/lex_bot/data/uploads"))
        os.makedirs(lex_bot_upload_dir, exist_ok=True)
        lex_bot_path = os.path.join(lex_bot_upload_dir, safe_name)
        with open(lex_bot_path, "wb") as f:
            f.write(docx_bytes)
            
        # Link to session cache if session_id is provided
        if session_id and get_session_cache:
            try:
                session_cache = get_session_cache()
                if session_cache:
                    session_cache.add_file_path(session_id, lex_bot_path)
            except Exception as cache_e:
                logger.warning(f"Failed to add path to session cache: {cache_e}")
                
        # Scan file for placeholders using python-docx and regex
        placeholders_detected = []
        try:
            from docx import Document
            doc = Document(output_path)
            # Scan paragraphs
            for p in doc.paragraphs:
                matches = PLACEHOLDER_REGEX.findall(p.text)
                for m in matches:
                    if m not in placeholders_detected:
                        placeholders_detected.append(m)
            # Scan tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            matches = PLACEHOLDER_REGEX.findall(p.text)
                            for m in matches:
                                if m not in placeholders_detected:
                                    placeholders_detected.append(m)
        except Exception as docx_e:
            logger.warning(f"Placeholder parsing failed: {docx_e}")
            
        document_key = hashlib.sha256(draft_id.encode("utf-8")).hexdigest()
        
        # Register draft metadata in PostgreSQL db via auth service
        try:
            async with httpx.AsyncClient() as client:
                await client.post(
                    f"{AUTH_SERVICE_URL.rstrip('/')}/internal/draft/register",
                    json={
                        "draft_id": draft_id,
                        "name": safe_name,
                        "filename": safe_name,
                        "document_key": document_key,
                        "created_by": user_id,
                        "variables_detected": placeholders_detected,
                        "status": "In progress"
                    },
                    timeout=10.0
                )
        except Exception as reg_err:
            logger.warning(f"Failed to register draft in DB: {reg_err}")

        params: Dict[str, Any] = {
            "document": {
                "fileType": "docx",
                "key": document_key,
                "title": safe_name,
                "url": f"http://drafter-service:8003/v2/draft/serve/{draft_id}/{safe_name}",
                "permissions": {"edit": True, "download": True, "print": True},
            },
            "documentType": "word",
            "editorConfig": {
                "callbackUrl": f"http://drafter-service:8003/v2/draft/callback/{draft_id}",
                "mode": "edit",
                "user": {
                    "id": user_id,
                    "name": f"User {user_id[:4]}" if len(user_id) >= 4 else f"User {user_id}"
                },
                "coauthoring": {
                    "mode": "fast",
                    "change": True
                },
                "customization": {
                    "forcesave": True,
                    "chat": True,
                    "uiTheme": "theme-light",
                    "logo": {
                        "image": "",
                        "imageDark": "",
                        "url": ""
                    }
                },
            },
        }
        params["token"] = _jwt_encode(params)
        params["variablesDetected"] = placeholders_detected
        params["documentKey"] = document_key
        params["filename"] = safe_name
        params["draftId"] = draft_id
        
        background_tasks.add_task(extract_and_cache_docx, lex_bot_path)
        
        return params
    except Exception as e:
        logger.exception("Draft upload failed.")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/v2/draft/forcesave")
async def onlyoffice_forcesave(request: ForceSaveRequest, authorization: Optional[str] = Header(default=None)):
    await verify_token(authorization)

    document_key = (request.document_key or "").strip()
    if not document_key:
        raise HTTPException(status_code=400, detail="document_key is required.")

    payload = {"c": "forcesave", "key": document_key}
    token = _jwt_encode(payload)
    command_payload = {"c": "forcesave", "key": document_key, "token": token}
    headers = {"Authorization": f"Bearer {token}"}

    try:
        async with httpx.AsyncClient() as client:
            resp = await client.post(
                f"{ONLYOFFICE_API_URL.rstrip('/')}/coauthoring/CommandService.ashx",
                json=command_payload,
                headers=headers,
                timeout=15.0,
            )
    except httpx.RequestError as e:
        raise HTTPException(status_code=502, detail=f"OnlyOffice CommandService request failed: {e}")

    if resp.status_code >= 400:
        raise HTTPException(status_code=502, detail=f"OnlyOffice CommandService error: {resp.status_code}")

    return {"ok": True}


@app.get("/v2/draft/config/{draft_id}")
async def get_draft_config(draft_id: str, authorization: Optional[str] = Header(default=None)):
    try:
        user_id = await verify_token(authorization)
        
        # 1. Verify access via auth service
        try:
            async with httpx.AsyncClient() as client:
                access_resp = await client.get(
                    f"{AUTH_SERVICE_URL.rstrip('/')}/internal/draft/verify_access/{draft_id}?user_id={user_id}",
                    timeout=5.0
                )
                access_resp.raise_for_status()
                access_data = access_resp.json()
        except Exception as exc:
            logger.error(f"Failed to contact auth service for access verification: {exc}")
            raise HTTPException(status_code=502, detail=f"Failed to contact auth service: {exc}")
             
        access_level = access_data.get("access_level", "none")
        if access_level == "none":
            raise HTTPException(status_code=403, detail="You do not have access to this draft.")
             
        # 2. Get draft metadata from auth service
        try:
            async with httpx.AsyncClient() as client:
                draft_resp = await client.get(
                    f"{AUTH_SERVICE_URL.rstrip('/')}/internal/draft/get/{draft_id}",
                    timeout=5.0
                )
                draft_resp.raise_for_status()
                draft_data = draft_resp.json()
        except Exception as exc:
            logger.error(f"Failed to fetch draft metadata from auth service: {exc}")
            raise HTTPException(status_code=502, detail=f"Failed to fetch draft info: {exc}")
             
        file_name = draft_data.get("filename")
        document_key = draft_data.get("documentKey")
         
        # 3. Build OnlyOffice config
        params: Dict[str, Any] = {
            "document": {
                "fileType": "docx",
                "key": document_key,
                "title": file_name,
                "url": f"http://drafter-service:8003/v2/draft/serve/{draft_id}/{file_name}",
                "permissions": {
                    "edit": access_level == "edit",
                    "download": True,
                    "print": True
                },
            },
            "documentType": "word",
            "editorConfig": {
                "callbackUrl": f"http://drafter-service:8003/v2/draft/callback/{draft_id}",
                "mode": "edit" if access_level == "edit" else "view",
                "user": {
                    "id": user_id,
                    "name": f"User {user_id[:4]}" if len(user_id) >= 4 else f"User {user_id}"
                },
                "coauthoring": {
                    "mode": "fast",
                    "change": True
                },
                "customization": {
                    "forcesave": True,
                    "chat": True,
                    "uiTheme": "theme-light",
                    "logo": {
                        "image": "",
                        "imageDark": "",
                        "url": ""
                    }
                },
                "plugins": {
                    "autostart": [
                        "asc.{43d1a84f-e274-4b53-a55e-3363f8db1f34}"
                    ],
                    "pluginsData": []
                }
            },
        }
        params["token"] = _jwt_encode(params)
        params["documentKey"] = document_key
        params["filename"] = file_name
        params["variablesDetected"] = draft_data.get("variablesDetected", [])
        params["draftId"] = draft_id
        params["status"] = draft_data.get("status", "In progress")
         
        return params
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.exception("Failed to load draft configuration.")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/v2/draft/callback/{draft_id}")
@app.post("/v2/draft/callback")
async def onlyoffice_callback(event: Dict[str, Any], draft_id: Optional[str] = None, authorization: Optional[str] = Header(default=None)):
    try:
        token: Optional[str] = None
        if isinstance(authorization, str) and authorization.strip():
            auth_val = authorization.strip()
            if auth_val.lower().startswith("bearer "):
                token = auth_val.split(" ", 1)[1].strip()
            else:
                token = auth_val
        if not token:
            payload_token = event.get("token")
            if isinstance(payload_token, str) and payload_token.strip():
                token = payload_token.strip()

        if not token:
            raise HTTPException(status_code=403, detail="Forbidden")

        try:
            decoded = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
            if isinstance(decoded, dict):
                if "payload" in decoded:
                    event = decoded["payload"]
                else:
                    event = decoded
        except jwt.PyJWTError:
            raise HTTPException(status_code=403, detail="Forbidden")

        status = event.get("status")
        if isinstance(status, str) and status.isdigit():
            status = int(status)

        if status == 4:
            return {"error": 0}

        if status in (2, 6):
            url = event.get("url") or event.get("fileUrl") or event.get("downloadUrl")
            if isinstance(url, dict):
                url = url.get("url")
            if not isinstance(url, str) or not url.strip():
                return {"error": 0}

            url = _normalize_download_url(url.strip())

            shared_storage_path = os.getenv("SHARED_STORAGE_PATH")
            if not shared_storage_path:
                return {"error": 0}

            file_name = None
            if draft_id:
                try:
                    async with httpx.AsyncClient() as client:
                        draft_resp = await client.get(
                            f"{AUTH_SERVICE_URL.rstrip('/')}/internal/draft/get/{draft_id}",
                            timeout=5.0
                        )
                        draft_resp.raise_for_status()
                        draft_data = draft_resp.json()
                        file_name = draft_data.get("filename")
                except Exception as exc:
                    logger.error(f"Failed to fetch draft metadata from auth service: {exc}")

            if not file_name:
                file_name = _safe_basename_from_url(url)

            if not file_name:
                return {"error": 0}

            if draft_id:
                target_dir = os.path.join(shared_storage_path, draft_id)
                os.makedirs(target_dir, exist_ok=True)
                target_path = os.path.join(target_dir, file_name)
            else:
                target_path = os.path.join(shared_storage_path, file_name)
                os.makedirs(shared_storage_path, exist_ok=True)

            async with httpx.AsyncClient(follow_redirects=True) as client:
                async with client.stream("GET", url, timeout=60.0) as resp:
                    resp.raise_for_status()
                    with open(target_path, "wb") as f:
                        async for chunk in resp.aiter_bytes():
                            if chunk:
                                f.write(chunk)
                                
            # Touch draft updated_at in DB
            if draft_id:
                try:
                    async with httpx.AsyncClient() as client:
                        await client.post(
                            f"{AUTH_SERVICE_URL.rstrip('/')}/internal/draft/touch/{draft_id}",
                            timeout=5.0
                        )
                except Exception as touch_err:
                    logger.warning(f"Failed to touch draft updated_at: {touch_err}")
    except Exception:
        logger.exception("OnlyOffice callback processing failed.")

    return {"error": 0}


if __name__ == "__main__":
    uvicorn.run("Drafter:app", host="0.0.0.0", port=8003, reload=True)
