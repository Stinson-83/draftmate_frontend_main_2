import os
import re
import json
import logging
import threading
from typing import List, Dict, Any, Optional
import httpx
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import google.generativeai as genai

logger = logging.getLogger(__name__)

# Initialize GenAI
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if GOOGLE_API_KEY:
    genai.configure(api_key=GOOGLE_API_KEY)

# Lazy-loaded EasyOCR reader to save startup latency
_ocr_reader = None

def get_ocr_reader():
    global _ocr_reader
    if _ocr_reader is None:
        try:
            import easyocr
            logger.info("Initializing EasyOCR reader...")
            _ocr_reader = easyocr.Reader(['en'])
        except ImportError:
            logger.warning("easyocr library not installed. Scanned PDF OCR fallback will be disabled.")
    return _ocr_reader


# ── Ingestion Helpers ────────────────────────────────────────────────

def extract_text_from_pdf(file_path: str, doc_id: str, auth_url: str) -> List[Dict[str, Any]]:
    """
    Extracts text page-by-page from a PDF. Runs OCR on blank pages.
    """
    pages_data = []
    doc = fitz.open(file_path)
    total_pages = len(doc)
    
    # Notify initial page count
    _update_doc_status(doc_id, "processing", 0, total_pages, auth_url)

    for page_idx in range(total_pages):
        page = doc[page_idx]
        page_num = page_idx + 1
        text = page.get_text().strip()

        # Fallback to OCR if page text is empty
        if not text:
            logger.info(f"Page {page_num} of {file_name_of(file_path)} is blank. Running OCR...")
            reader = get_ocr_reader()
            if reader:
                try:
                    # Render page to image bytes for OCR
                    pix = page.get_pixmap(dpi=150)
                    img_data = pix.tobytes("png")
                    ocr_results = reader.readtext(img_data)
                    text = " ".join([res[1] for res in ocr_results]).strip()
                except Exception as ocr_err:
                    logger.error(f"OCR failed on page {page_num}: {ocr_err}")

        pages_data.append({
            "page_number": page_num,
            "text": text
        })
        
        # Notify progress
        _update_doc_status(doc_id, "processing", page_num, total_pages, auth_url)

    doc.close()
    return pages_data


def extract_text_from_docx(file_path: str, doc_id: str, auth_url: str) -> List[Dict[str, Any]]:
    """
    Parses DOCX files and maps text segments to estimated pages.
    """
    doc = DocxDocument(file_path)
    text_runs = []
    for p in doc.paragraphs:
        if p.text.strip():
            text_runs.append(p.text)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        text_runs.append(p.text)
                        
    full_text = "\n".join(text_runs)
    
    # Simple word count approximation for pages (approx 350 words per page)
    words = full_text.split()
    total_words = len(words)
    words_per_page = 350
    total_pages = max(1, (total_words + words_per_page - 1) // words_per_page)
    
    _update_doc_status(doc_id, "processing", 0, total_pages, auth_url)
    
    pages_data = []
    for page_idx in range(total_pages):
        start_w = page_idx * words_per_page
        end_w = min(total_words, (page_idx + 1) * words_per_page)
        page_text = " ".join(words[start_w:end_w])
        pages_data.append({
            "page_number": page_idx + 1,
            "text": page_text
        })
        _update_doc_status(doc_id, "processing", page_idx + 1, total_pages, auth_url)
        
    return pages_data


def extract_text_from_txt(file_path: str, doc_id: str, auth_url: str) -> List[Dict[str, Any]]:
    """
    Reads plain text files.
    """
    _update_doc_status(doc_id, "processing", 0, 1, auth_url)
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    
    _update_doc_status(doc_id, "processing", 1, 1, auth_url)
    return [{"page_number": 1, "text": text}]


def file_name_of(file_path: str) -> str:
    return os.path.basename(file_path)


def _update_doc_status(doc_id: str, status: str, pages_processed: int, total_pages: int, auth_url: str):
    try:
        with httpx.Client() as client:
            client.post(
                f"{auth_url.rstrip('/')}/internal/chronology/document/update",
                json={
                    "doc_id": doc_id,
                    "status": status,
                    "pages_processed": pages_processed,
                    "total_pages": total_pages
                },
                timeout=5.0
            )
    except Exception as e:
        logger.error(f"Failed to update document status: {e}")


# ── Asynchronous Background Parser Runner ─────────────────────────────

def process_document_async(file_path: str, doc_id: str, auth_url: str, upload_dir: str):
    """
    Runs parser and writes the page index context to a JSON file.
    """
    def run():
        try:
            ext = os.path.splitext(file_path)[1].lower()
            if ext == ".pdf":
                pages = extract_text_from_pdf(file_path, doc_id, auth_url)
            elif ext == ".docx":
                pages = extract_text_from_docx(file_path, doc_id, auth_url)
            else:
                pages = extract_text_from_txt(file_path, doc_id, auth_url)
                
            # Write page contents to cache file
            cache_path = os.path.join(upload_dir, f"{doc_id}_parsed.json")
            with open(cache_path, "w", encoding="utf-8") as f:
                json.dump(pages, f, ensure_ascii=False, indent=2)
                
            # Mark completion
            _update_doc_status(doc_id, "done", len(pages), len(pages), auth_url)
            logger.info(f"Successfully processed document: {doc_id}")
        except Exception as err:
            logger.error(f"Error parsing document {doc_id}: {err}")
            _update_doc_status(doc_id, "failed", 0, 0, auth_url)
            
    threading.Thread(target=run, daemon=True).start()


# ── AI Chronology Builder Engine ─────────────────────────────────────

AI_EXTRACTION_PROMPT = """
You are an expert Legal Counsel and Master Timeline Auditor.
Your task is to review the following text from a case document and extract a list of material events with exact dates and page-level attribution.

GUIDELINES:
1. Extract dates, date types (exact/inferred/approximate), factual descriptions of actions/events, actors involved, and reference paragraphs.
2. Keep descriptions objective and neutral. Avoid legal conclusions (e.g. write "Party A terminated the agreement" instead of "Party A unlawfully terminated").
3. Preserve the exact source text snippet in the output.
4. If a date is inferred, set date_type to "inferred" and explain your inference in source_text.
5. Format your output strictly as a JSON object matching this schema:
{
  "events": [
    {
      "event_date": "YYYY-MM-DD",
      "date_type": "exact" | "inferred" | "approximate",
      "event_description": "Objective description",
      "actors": ["Actor1", "Actor2"],
      "source_text": "Original quote from passage...",
      "confidence": 0.95
    }
  ]
}
"""

def extract_events_from_page(page_text: str, page_num: int, doc_name: str) -> List[Dict[str, Any]]:
    """
    Calls Gemini to parse chronology events from a single page's text context.
    """
    if not page_text.strip():
        return []

    try:
        model = genai.GenerativeModel(
            model_name="gemini-2.5-flash",
            system_instruction=AI_EXTRACTION_PROMPT,
            generation_config=genai.GenerationConfig(
                temperature=0.1,
                response_mime_type="application/json",
                max_output_tokens=2048
            )
        )
        prompt_content = f"DOCUMENT SOURCE: {doc_name}\nPAGE NUMBER: {page_num}\nTEXT CONTENT:\n{page_text}"
        response = model.generate_content(prompt_content)
        
        # Parse output
        raw_text = response.text.strip()
        if raw_text.startswith("```"):
            lines = raw_text.splitlines()
            if lines[0].startswith("```"):
                lines = lines[1:]
            if lines[-1].startswith("```"):
                lines = lines[:-1]
            raw_text = "\n".join(lines).strip()

        try:
            payload = json.loads(raw_text)
        except Exception as json_err:
            logger.warning(f"Standard JSON decode failed on page {page_num}: {json_err}. Trying AST fallback.")
            try:
                import ast
                payload = ast.literal_eval(raw_text)
                if not isinstance(payload, dict):
                    raise ValueError("Parsed object is not a dictionary")
            except Exception as ast_err:
                logger.error(f"Both JSON and AST parsing failed on page {page_num}. Raw response: {raw_text}")
                raise json_err

        events = payload.get("events", [])
        
        # Append source metadata
        for ev in events:
            ev["source_document"] = doc_name
            ev["source_page"] = page_num
            ev["is_conflict"] = False
            ev["conflict_details"] = []
            ev["status"] = "pending"
            
        return events
    except Exception as e:
        logger.error(f"Failed to extract AI events on page {page_num} of {doc_name}: {e}")
        return []


def run_chronology_engine(raw_events: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Performs date normalization, duplicate detection (merging identical events),
    and conflict detection across the raw event list.
    """
    # 1. Deduplication and Merge Grouping
    # Group events that occur on the exact same date and have similar descriptions
    merged_events: List[Dict[str, Any]] = []
    
    for raw in raw_events:
        event_date = raw.get("event_date", "9999-99-99")
        desc = raw.get("event_description", "").lower().strip()
        
        # Check if we already have a similar event on this date
        matched = False
        for merged in merged_events:
            if merged["event_date"] == event_date:
                # Basic similarity check: shared actors or high keyword overlap
                exist_desc = merged["event_description"].lower()
                # If they share significant words or match semantically
                intersection = set(desc.split()) & set(exist_desc.split())
                if len(intersection) >= 3 or desc in exist_desc or exist_desc in desc:
                    # Append source document/page to citations
                    # Save a list of citations instead of just a string
                    if "citations" not in merged:
                        merged["citations"] = [{
                            "source_document": merged["source_document"],
                            "source_page": merged["source_page"],
                            "source_text": merged.get("source_text", "")
                        }]
                    
                    merged["citations"].append({
                        "source_document": raw["source_document"],
                        "source_page": raw["source_page"],
                        "source_text": raw.get("source_text", "")
                    })
                    
                    # Update actor list
                    merged["actors"] = list(set(merged["actors"] + raw.get("actors", [])))
                    matched = True
                    break
                    
        if not matched:
            # First instance, create event structure
            new_evt = dict(raw)
            new_evt["citations"] = [{
                "source_document": raw["source_document"],
                "source_page": raw["source_page"],
                "source_text": raw.get("source_text", "")
            }]
            merged_events.append(new_evt)

    # 2. Conflict Detection (Differing dates for same description/actors)
    # If the same event (e.g. "termination of agreement") has different dates in different documents
    for i in range(len(merged_events)):
        for j in range(i + 1, len(merged_events)):
            evt1 = merged_events[i]
            evt2 = merged_events[j]
            
            if evt1["event_date"] != evt2["event_date"]:
                desc1 = evt1["event_description"].lower()
                desc2 = evt2["event_description"].lower()
                
                # Check if descriptions are very similar
                intersection = set(desc1.split()) & set(desc2.split())
                if len(intersection) >= 4 or desc1 in desc2 or desc2 in desc1:
                    # Mark conflict
                    evt1["is_conflict"] = True
                    evt2["is_conflict"] = True
                    
                    conflict_info_1 = {
                        "date": evt1["event_date"],
                        "source": f"{evt1['source_document']} p.{evt1['source_page']}"
                    }
                    conflict_info_2 = {
                        "date": evt2["event_date"],
                        "source": f"{evt2['source_document']} p.{evt2['source_page']}"
                    }
                    
                    evt1["conflict_details"].append(conflict_info_2)
                    evt2["conflict_details"].append(conflict_info_1)

    return merged_events
