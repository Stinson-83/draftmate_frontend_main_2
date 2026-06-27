"""DOCX reconstruction helpers for translated blocks."""

from __future__ import annotations

from collections import defaultdict
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from backend.translator.extractors.models import Block


def _apply_run_style(run, block: Block) -> None:
    style = block.style or {}
    if style.get("font_name"):
        run.font.name = str(style["font_name"])
    if style.get("font_size"):
        run.font.size = Pt(float(style["font_size"]))
    if style.get("bold") is not None:
        run.bold = bool(style["bold"])
    if style.get("italic") is not None:
        run.italic = bool(style["italic"])


def _add_text_paragraph(document: Document, block: Block) -> None:
    paragraph = document.add_paragraph()
    style = block.style or {}

    if block.bounding_box != (0.0, 0.0, 0.0, 0.0):
        paragraph.paragraph_format.left_indent = Pt(float(block.bounding_box[0]))
        paragraph.paragraph_format.space_before = Pt(max(float(block.bounding_box[1]), 0.0))

    alignment = style.get("alignment")
    if alignment == "CENTER":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "RIGHT":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == "JUSTIFY":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = paragraph.add_run(block.text)
    _apply_run_style(run, block)


def _add_table(document: Document, table_blocks: Sequence[Block]) -> None:
    rows: dict[int, dict[int, Block]] = defaultdict(dict)
    for block in table_blocks:
        row_index = int(block.style.get("row_index", 1))
        column_index = int(block.style.get("column_index", 1))
        rows[row_index][column_index] = block

    max_columns = max((max(columns) for columns in rows.values()), default=0)
    if max_columns == 0:
        return

    table = document.add_table(rows=len(rows), cols=max_columns)
    table.style = "Table Grid"

    for row_offset, row_index in enumerate(sorted(rows)):
        for column_offset in range(1, max_columns + 1):
            cell = table.cell(row_offset, column_offset - 1)
            block = rows[row_index].get(column_offset)
            cell.text = block.text if block else ""


def rebuild_docx_document(blocks: Sequence[Block], output_path: str | Path) -> Path:
    document = Document()
    emitted_tables: set[int] = set()
    table_groups: dict[int, list[Block]] = defaultdict(list)

    for block in blocks:
        if block.type == "table_cell" and "table_index" in (block.style or {}):
            table_groups[int(block.style["table_index"])] .append(block)

    for block in blocks:
        if block.type == "table_cell" and "table_index" in (block.style or {}):
            table_index = int(block.style["table_index"])
            if table_index not in emitted_tables:
                _add_table(document, table_groups[table_index])
                emitted_tables.add(table_index)
            continue

        if block.type.startswith("heading"):
            level = int((block.style or {}).get("level", 1) or 1)
            heading = document.add_heading(level=max(1, min(level, 9)))
            run = heading.add_run(block.text)
            _apply_run_style(run, block)
        else:
            _add_text_paragraph(document, block)

    output_path = Path(output_path)
    document.save(str(output_path))
    return output_path
