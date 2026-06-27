"""DOCX extraction helpers using python-docx."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from .models import Block


def _block_style(paragraph) -> dict[str, Any]:
    style: dict[str, Any] = {
        "style_name": getattr(getattr(paragraph, "style", None), "name", None),
        "alignment": getattr(paragraph.alignment, "name", None) if paragraph.alignment else None,
    }

    for run in paragraph.runs:
        if run.text.strip():
            style.update(
                {
                    "font_name": getattr(run.font, "name", None),
                    "font_size": float(run.font.size.pt) if getattr(run.font, "size", None) else None,
                    "bold": run.bold,
                    "italic": run.italic,
                }
            )
            break

    return style


def extract_docx_blocks(docx_path: str | Path) -> list[Block]:
    document = Document(docx_path)
    blocks: list[Block] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        blocks.append(
            Block(
                id=len(blocks) + 1,
                type="paragraph",
                text=text,
                bounding_box=(0.0, 0.0, 0.0, 0.0),
                style=_block_style(paragraph),
            )
        )

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for column_index, cell in enumerate(row.cells, start=1):
                text = "\n".join(paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip())
                if not text:
                    continue

                blocks.append(
                    Block(
                        id=len(blocks) + 1,
                        type="table_cell",
                        text=text,
                        bounding_box=(0.0, 0.0, 0.0, 0.0),
                        style={
                            "table_index": table_index,
                            "row_index": row_index,
                            "column_index": column_index,
                        },
                    )
                )

    return blocks
