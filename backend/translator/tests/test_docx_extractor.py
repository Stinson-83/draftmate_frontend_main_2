from pathlib import Path

from docx import Document

from backend.translator.extractors import extract_docx_blocks


def _create_sample_docx(docx_path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Intro paragraph")
    run.bold = True
    document.add_paragraph("Second paragraph")

    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table cell text"

    document.save(docx_path)


def test_extract_docx_blocks(tmp_path: Path) -> None:
    docx_path = tmp_path / "sample.docx"
    _create_sample_docx(docx_path)

    blocks = extract_docx_blocks(docx_path)

    assert [block.text for block in blocks[:3]] == ["Intro paragraph", "Second paragraph", "Table cell text"]
    assert blocks[0].type == "paragraph"
    assert blocks[2].type == "table_cell"
    assert blocks[0].style["bold"] is True
    assert blocks[0].bounding_box == (0.0, 0.0, 0.0, 0.0)
