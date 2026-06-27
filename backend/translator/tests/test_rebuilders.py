from pathlib import Path

import pdfplumber
from docx import Document

from backend.translator.extractors.models import Block
from backend.translator.rebuilders import build_html_document, rebuild_docx_document, rebuild_html_document, rebuild_pdf_document


def _sample_blocks() -> list[Block]:
    return [
        Block(
            id=1,
            type="heading",
            text="Translated Title",
            bounding_box=(24.0, 24.0, 260.0, 50.0),
            style={"font_name": "Helvetica", "font_size": 16, "bold": True, "page_number": 1, "level": 1},
        ),
        Block(
            id=2,
            type="paragraph",
            text="First translated paragraph.",
            bounding_box=(24.0, 70.0, 300.0, 100.0),
            style={"font_name": "Helvetica", "font_size": 12, "page_number": 1},
        ),
        Block(
            id=3,
            type="table_cell",
            text="Cell A1",
            bounding_box=(0.0, 0.0, 0.0, 0.0),
            style={"table_index": 1, "row_index": 1, "column_index": 1},
        ),
        Block(
            id=4,
            type="table_cell",
            text="Cell A2",
            bounding_box=(0.0, 0.0, 0.0, 0.0),
            style={"table_index": 1, "row_index": 1, "column_index": 2},
        ),
    ]


def test_build_html_document_contains_reconstructed_blocks() -> None:
    html = build_html_document(_sample_blocks())

    assert "Translated Title" in html
    assert "First translated paragraph." in html
    assert "translation-table" in html
    assert "position: absolute" in html


def test_rebuild_html_document_writes_file(tmp_path: Path) -> None:
    output_path = tmp_path / "document.html"

    rebuilt_path = rebuild_html_document(_sample_blocks(), output_path)

    assert rebuilt_path == output_path
    html = output_path.read_text(encoding="utf-8")
    assert "Cell A1" in html
    assert "Translated Title" in html


def test_rebuild_docx_document_writes_structure(tmp_path: Path) -> None:
    output_path = tmp_path / "document.docx"

    rebuilt_path = rebuild_docx_document(_sample_blocks(), output_path)

    assert rebuilt_path == output_path
    document = Document(str(output_path))
    paragraph_texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    assert "Translated Title" in paragraph_texts[0]
    assert "First translated paragraph." in paragraph_texts[1]
    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 0).text == "Cell A1"
    assert document.tables[0].cell(0, 1).text == "Cell A2"


def test_rebuild_pdf_document_writes_text(tmp_path: Path) -> None:
    output_path = tmp_path / "document.pdf"

    rebuilt_path = rebuild_pdf_document(_sample_blocks(), output_path)

    assert rebuilt_path == output_path
    assert output_path.exists()

    with pdfplumber.open(str(output_path)) as pdf:
        extracted_text = "\n".join(page.extract_text() or "" for page in pdf.pages)

    assert "Translated Title" in extracted_text
    assert "First translated paragraph." in extracted_text
